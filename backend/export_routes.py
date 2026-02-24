from flask import Blueprint, request, jsonify, send_file, g
from io import BytesIO
from openpyxl import Workbook
from docx import Document as WordDoc
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from models import UploadedFile
from models import db, Document
from utils.jwt_utils import jwt_required
import re

export_bp = Blueprint("export", __name__, url_prefix="/export")

# Marathi → English mapping
DOCNAME_MAP = {
    "Agreement Relating to Deposit of Title Deeds,Pawn": "Notice of Intimation",
    "Notice of Intimation": "Notice of Intimation",

    "डिपॉझिट": "Mortgage Deed",
    "गहाणखत": "Mortgage Deed",
    "बक्षीसपत्र": "Mortgage Deed",
    "जामाखत": "Mortgage Deed",
    "फर्दर चार्ज": "Mortgage Deed",
    "रीकन्वेअन्स": "Reconveyance",
    "घोषणापत्र": "Declaration Deed",
    "हक्कसोड पत्र": "Release Deed",
    "खरेदीखत": "Sale Deed",
    "बक्षीसपत्र": "Gift Deed",
    "भाडेपट्टा": "Lease Deed",
    "करारनामा": "Agreement",
    "परिमोचनपत्र":" Relees Deed",
    "संमती पत्र": "Consent Deed",
    "लिव्ह अँड लायसेन्स": "Leave and Licence",
    "सुधार पत्र": "Correction Deed",
    "ताबा पत्र": "Possession Letter",
    "मान्यतापत्र":"Consent Deed",
    "1236-अ-लिव्ह अॅन्ड लायसन्सेस":"Leave and Licence",
    "वाडवाणी": "Partition Deed",
    "राखणी नामा": "Partition Deed",
    "चुक दुरतस्ती":"Correction Deed",
    "रद्दपत्र": "Cancellation Deed",
    "अधिकार हस्तांतरणाची विक्री": "Sale Deed",
    "अभिहस्तांतरणपत्र": "Gift Deed",
    "अपार्टमेंट डीड":"Deed of Apartment",
    "डीड ऑफ एक्स्चेंज": "Deed of Exchange",
    "नोटीस ऑफ लिस पेंडेन्स": "Notice of Lease Pendency",
}

def clean_filename(name: str) -> str:
    # remove illegal Windows filename characters
    name = name.replace(" ", "_")
    return re.sub(r'[\\/*?:"<>|]', "", name)
# ====================================================================
# EXPORT SELECTED → EXCEL
# ====================================================================
@export_bp.route("/selected/excel", methods=["POST"])
@jwt_required
def export_selected_excel():
    data = request.get_json()
    ids = [e.get("id") for e in data.get("entries", [])]

    if not ids:
        return jsonify({"error": "No entries selected"}), 400

    docs = Document.query.filter(
        Document.id.in_(ids),
        Document.user_id == g.current_user.id
    ).all()

    wb = Workbook()
    ws = wb.active
    ws.append([
        "ID", "Doc No", "Doc Name", "Purchaser", "Seller",
        "Registration Date", "Area", "Consideration"
    ])

    for d in docs:
        ws.append([
            d.id,
            d.docno,
            d.docname,
            d.purchasername,
            d.sellername,
            d.registrationdate,
            d.areaname,
            d.consideration_amt
        ])

    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)

    return send_file(
        stream,
        as_attachment=True,
        download_name="selected_entries.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


# ====================================================================
# EXPORT SELECTED → WORD (Formatted Title + Table)
# ====================================================================
@export_bp.route("/selected/word", methods=["POST"])
@jwt_required
def export_selected_word():

    data = request.get_json()

    ids = [e.get("id") for e in data.get("entries", [])]
    table_name = data.get("table_name", "").strip()   # <<< NEW

    if not ids:
        return jsonify({"error": "No entries selected"}), 400

    # ---------- GET FILE METADATA ----------
    uploaded_file = UploadedFile.query.filter_by(
        table_name=table_name,
        user_id=g.current_user.id
    ).first()

    if not uploaded_file:
        return jsonify({"error": "File metadata not found"}), 404

    # ---------- FETCH DOCUMENTS ----------
    docs = Document.query.filter(
        Document.id.in_(ids),
        Document.user_id == g.current_user.id
    ).all()

    if not docs:
        return jsonify({"error": "No matching documents"}), 400

    docx = WordDoc()

    for d in docs:
        # English doc name conversion
        eng_docname = DOCNAME_MAP.get(d.docname.strip(), d.docname)

        # Extract year
        year = ""
        if d.registrationdate:
            year = d.registrationdate.split("-")[0]

        # Convert SRO if Marathi
        sro_eng = ""
        sro_lower = d.sroname.lower()
        if "हवेली" in d.sroname or "haveli" in sro_lower:
            num = "".join([c for c in d.sroname if c.isdigit()])
            sro_eng = f"HVL {num}"

        title = f"{year} – {eng_docname} dated {d.registrationdate} (Reg. No {sro_eng} {d.docno}/{year})"

        # Title
        p = docx.add_paragraph(title)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(13)

        # Table
        table = docx.add_table(rows=1, cols=3)
        table.autofit = True

        hdr = table.rows[0].cells
        hdr[0].text = d.sellername or ""
        hdr[1].text = d.purchasername or ""
        hdr[2].text = d.propertydescription or ""

        docx.add_paragraph("\n")

    buffer = BytesIO()
    docx.save(buffer)
    buffer.seek(0)

    # ---------- CREATE DOWNLOAD FILE NAME ----------
    filename = f"{uploaded_file.table_name}.docx"
    filename = clean_filename(filename)

    response = send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response